from pathlib import Path
from docx import Document as DocxDocument
from backend.rag_engine.models.document import Document


class DocxLoader:
    def load(self, file_path: str) -> list[Document]:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"{path} does not exist.")

        doc = DocxDocument(path)

        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        content = "\n".join(paragraphs)

        return [
            Document(
                content=content,
                source=path.name,
                file_type="docx",
                metadata={
                    "paragraphs": len(paragraphs),
                    "file_size": path.stat().st_size,
                },
            )
        ]
