from pathlib import Path
from rag_engine.models.document import Document
from rag_engine.ingestion.base_loader import BaseLoader
from core.logger import logger


class DocxLoader(BaseLoader):
    def load(self, file_path: str, progress=None) -> list[Document]:
        path = self.validate(file_path)
        if progress:
            progress.upload(
                "loading_docx",
                "Reading DOCX",
                20,
            )

        # lazy
        from docx import Document as DocxDocument

        logger.info(f"Loading Docx: {path.name}")
        file_size = path.stat().st_size
        try:
            doc = DocxDocument(path)

            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

            content = "\n".join(paragraphs)

            documents = [
                Document(
                    content=content,
                    source=path.name,
                    file_type="docx",
                    metadata={
                        "paragraphs": len(paragraphs),
                        "file_size": file_size,
                    },
                )
            ]
        except Exception:
            logger.exception(f"Failed loading Docx: {path.name}")
            raise

        if progress:
            progress.upload(
                "pdf_loaded",
                f"Loaded {len(documents)} page.",
                25,
            )
        logger.info(f"Loaded {len(documents)} pages from {path.name}")
        return documents
